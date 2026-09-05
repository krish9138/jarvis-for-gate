import os
from pathlib import Path
from typing import List, Dict, Any

def extract_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF file.
    Returns: [{'page_number': 1, 'text': '...'}, ...]
    """
    pages = []
    
    # Try pypdf first
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({
                "page_number": idx + 1,
                "text": text.strip()
            })
        return pages
    except Exception as e_pypdf:
        pass

    # Fallback to PyPDF2
    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": idx + 1,
                    "text": text.strip()
                })
        return pages
    except Exception as e_pypdf2:
        raise RuntimeError(f"Failed to extract PDF text from {file_path}: {e_pypdf2}")

def extract_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a .docx file paragraph-by-paragraph or by section.
    Returns: [{'page_number': 1, 'text': '...'}, ...]
    """
    try:
        import docx
        doc = docx.Document(file_path)
        
        paragraphs = []
        current_section = []
        section_idx = 1
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            if p.style.name.startswith("Heading") and current_section:
                paragraphs.append({
                    "page_number": section_idx,
                    "text": "\n".join(current_section)
                })
                current_section = [text]
                section_idx += 1
            else:
                current_section.append(text)
                
        if current_section:
            paragraphs.append({
                "page_number": section_idx,
                "text": "\n".join(current_section)
            })
            
        # If no sections were separated, bundle entire text
        if not paragraphs:
            full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            paragraphs = [{"page_number": 1, "text": full_text}]
            
        return paragraphs
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")

def extract_from_txt(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a plain text or markdown file.
    Supports utf-8, latin-1, and utf-16 encodings.
    """
    encodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
    content = ""
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                content = f.read()
            break
        except Exception:
            continue
            
    if not content:
        raise RuntimeError(f"Could not read text file with standard encodings: {file_path}")

    # Split by double formfeed or markdown large headings if present, or return as single section
    sections = [s.strip() for s in content.split("\n\n---\n\n") if s.strip()]
    if len(sections) > 1:
        return [{"page_number": idx + 1, "text": sec} for idx, sec in enumerate(sections)]
    else:
        return [{"page_number": 1, "text": content.strip()}]

def extract_document_text(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    """
    Main dispatcher for document text extraction.
    Returns: List of page dictionaries [{'page_number': int, 'text': str}]
    """
    ext = file_type.lower().strip().replace(".", "")
    if ext == "pdf":
        return extract_from_pdf(file_path)
    elif ext in ["docx", "doc"]:
        return extract_from_docx(file_path)
    elif ext in ["txt", "md", "markdown"]:
        return extract_from_txt(file_path)
    else:
        # Default text attempt
        return extract_from_txt(file_path)
